"""
Генерация отчёта по лабораторной работе №3.
Запуск: python generate_report.py
Результат: report_lab2.docx
После генерации вставьте скриншоты графиков из ноутбука вместо [ВСТАВИТЬ РИСУНОК].
"""
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# === Настройка страницы ===
sec = doc.sections[0]
sec.top_margin = Emu(914400)
sec.bottom_margin = Emu(914400)
sec.left_margin = Emu(1143000)
sec.right_margin = Emu(1143000)

style = doc.styles['Normal']
style.font.size = Pt(14)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15


def add_centered(text, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.bold = bold
    return p

def add_normal(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.bold = bold
    return p

def add_mixed(parts):
    """parts: list of (text, bold) tuples"""
    p = doc.add_paragraph()
    for text, bold in parts:
        r = p.add_run(text)
        r.font.size = Pt(14)
        r.bold = bold
    return p

def add_code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = 'Courier New'
    return p

def add_empty(n=1):
    for _ in range(n):
        doc.add_paragraph()

def add_figure_placeholder(caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('[ВСТАВИТЬ РИСУНОК]')
    r.font.size = Pt(12)
    r.italic = True
    p2 = doc.add_paragraph()
    r2 = p2.add_run(caption)
    r2.font.size = Pt(12)
    return p2

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for j, row in enumerate(rows):
        for i, val in enumerate(row):
            t.rows[j + 1].cells[i].text = str(val)
    return t


# ============================
# ТИТУЛЬНАЯ СТРАНИЦА
# ============================
add_centered('ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ ОБРАЗОВАТЕЛЬНОЕ')
add_centered('УЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ')
add_centered('РОССИЙСКИЙ УНИВЕРСИТЕТ ДРУЖБЫ НАРОДОВ')
add_centered('ИМЕНИ ПАТРИСА ЛУМУМБЫ')
add_empty(8)
add_centered('Отчет')
add_centered('о выполнении лабораторной работы № 3')
add_centered('"Многослойные сети. Алгоритм обратного распространения ошибки"')
add_centered('по дисциплине "Нейронные сети"')
add_empty()
add_centered('Вариант задания № 12')
add_empty(3)
add_normal('Выполнил студент группы ЗФИбд-01-24')
add_normal('Князев А. М. __________')
add_empty()
add_normal('Проверил и принял')
add_normal('Леонов С. С. __________')
add_empty()
add_normal('с оценкой __________')
add_empty(4)
add_centered('Москва, 2026')
doc.add_page_break()

# ============================
# ЦЕЛЬ РАБОТЫ
# ============================
add_mixed([
    ('Цель работы: ', True),
    ('исследование свойств многослойной нейронной сети прямого распространения '
     'и алгоритмов ее обучения, применение сети в задачах классификации '
     'и аппроксимации функции.', False),
])
add_empty()

add_normal('Основные этапы работы:', bold=True)
add_normal('1. Использовать многослойную нейронную сеть для классификации точек '
           'в случае, когда классы не являются линейно разделимыми.')
add_normal('2. Использовать многослойную нейронную сеть для аппроксимации функции. '
           'Произвести обучение с помощью метода первого порядка (traingdx).')
add_normal('3. Использовать многослойную нейронную сеть для аппроксимации функции. '
           'Произвести обучение с помощью метода второго порядка (trainbfg).')
add_empty()

add_mixed([
    ('Оборудование: ', True),
    ('персональный компьютер. ', False),
    ('Программное обеспечение: ', True),
    ('Python 3.x, Jupyter Notebook, библиотеки PyTorch, NumPy и Matplotlib.', False),
])
add_empty()

# ============================
# СЦЕНАРИЙ ВЫПОЛНЕНИЯ
# ============================
add_normal('Сценарий выполнения работы', bold=True)
add_empty()

# --- ЭТАП 1 ---
add_normal('Этап 1. Классификация линейно неразделимых классов', bold=True)
add_empty()

add_normal('Заданы 3 линейно неразделимых класса. Точки, принадлежащие одному классу, '
           'лежат на алгебраической линии. Параметры кривых приведены в таблице 1.')
add_empty()

add_normal('Таблица 1 — Параметры кривых для генерации точек')
add_table(
    ['Класс', 'Тип кривой', 'Параметры', 'Число точек'],
    [
        ['1', 'Эллипс (окружность)', 'a = b = 0.3, x₀ = 0, y₀ = 0', '60'],
        ['2', 'Эллипс (окружность)', 'a = b = 0.7, x₀ = 0, y₀ = 0', '100'],
        ['3', 'Парабола y² = 2px', 'p = 1, x₀ = −0.8, y₀ = 0', '120'],
    ]
)
add_empty()

add_normal('Для генерации точек использовались параметрические уравнения. '
           'Для эллипсов: x = a·cos(t), y = b·sin(t), t ∈ [0, 2π]. '
           'Для параболы: x = t²/2 + x₀, y = t, t ∈ [−2, 2]. '
           'Из полученных множеств случайным образом (randperm) выбраны '
           '60, 100 и 120 точек соответственно.')
add_empty()

add_normal('Множество точек каждого класса разделено на обучающее (70%), '
           'контрольное (20%) и тестовое (10%) подмножества с помощью функции dividerand. '
           'Соответствующие подмножества объединены в общую обучающую выборку.')
add_empty()

add_figure_placeholder('Рис. 1 — Исходные множества точек для трёх классов')
add_empty()

add_normal('Таблица 2 — Структура сети (Этап 1)')
add_table(
    ['Параметр', 'Значение'],
    [
        ['Тип сети', 'feedforwardnet (многослойная прямого распространения)'],
        ['Архитектура', '[2] → [20, tansig] → [3, sigmoid]'],
        ['Алгоритм обучения', 'RProp (trainrp)'],
        ['Число эпох (epochs)', '1500'],
        ['max_fail', '1500'],
        ['goal', '1e-5'],
        ['Число параметров', '103'],
    ]
)
add_empty()

add_normal('Сеть создана с 20 нейронами в скрытом слое с функцией активации tansig (tanh) '
           'и 3 нейронами в выходном слое с функцией активации sigmoid. '
           'Входное множество лежит в диапазоне [−1.2, 1.2], '
           'выходное — в диапазоне [0, 1] по каждой из координат.')
add_empty()

add_normal('Весовые коэффициенты инициализированы с помощью функции по умолчанию (init). '
           'Обучение выполнено алгоритмом RProp.')
add_empty()

add_figure_placeholder('Рис. 2 — Performance (график обучения, Этап 1)')
add_empty()

add_normal('Результаты классификации:', bold=True)
add_normal('Выход сети преобразован по правилу: oij = 1 если aij > 0.5, иначе 0.')
add_empty()

add_normal('Таблица 3 — Результаты классификации')
add_table(
    ['Подмножество', 'Правильно', 'Всего', 'Точность'],
    [
        ['Обучающее', '___', '___', '___%'],
        ['Контрольное', '___', '___', '___%'],
        ['Тестовое', '___', '___', '___%'],
    ]
)
add_empty()

add_normal('Для классификации области [−1.2, 1.2] × [−1.2, 1.2] задана сетка '
           'с шагом h = 0.025. Выход сети для каждой точки определяет принадлежность '
           'к трём классам. Компоненты выходного вектора задают интенсивность цветов RGB: '
           '(1,0,0) — красный (класс 1), (0,1,0) — зелёный (класс 2), '
           '(0,0,1) — синий (класс 3).')
add_empty()

add_figure_placeholder('Рис. 3 — Классификация области (RGB = выход сети)')

doc.add_page_break()

# --- ЭТАП 2 ---
add_normal('Этап 2. Аппроксимация функции (метод первого порядка)', bold=True)
add_empty()

add_normal('Задан обучающий набор {t, x(t)}, где x = cos(t² − 2t + 3), '
           't ∈ [0, 5], шаг h = 0.02. Всего 251 точка.')
add_empty()

add_normal('С конца временной последовательности выделены 10% отсчётов '
           'на контрольное подмножество. Тестовое подмножество оставлено пустым.')
add_empty()

add_figure_placeholder('Рис. 4 — Целевая функция x = cos(t² − 2t + 3)')
add_empty()

add_normal('Таблица 4 — Структура сети (Этап 2)')
add_table(
    ['Параметр', 'Значение'],
    [
        ['Архитектура', '[1] → [10, tansig] → [1, purelin]'],
        ['Алгоритм обучения', 'traingdx (SGD + момент + адаптивный lr)'],
        ['lr', '0.05'],
        ['mc (момент)', '0.9'],
        ['lr_inc', '1.05'],
        ['Число эпох (epochs)', '2000'],
        ['max_fail', '600'],
        ['goal', '1e-8'],
    ]
)
add_empty()

add_normal('Метод traingdx реализует градиентный спуск с моментом и адаптивным шагом. '
           'При уменьшении ошибки скорость обучения увеличивается в lr_inc раз. '
           'При значительном увеличении ошибки (более 4%) веса откатываются, '
           'а скорость обучения уменьшается.')
add_empty()

add_normal('Обучение проводилось несколько раз с различными начальными весами. '
           'Выбран лучший результат по минимальной MSE на обучающем подмножестве.')
add_empty()

add_normal('Весовые коэффициенты и смещения (после обучения):', bold=True)
add_normal('Скрытый слой: W1 = [___], b1 = [___]')
add_normal('Выходной слой: W2 = [___], b2 = [___]')
add_empty()

add_figure_placeholder('Рис. 5 — Performance (Этап 2 — traingdx)')
add_empty()

add_normal('Таблица 5 — Показатели качества (Этап 2)')
add_table(
    ['Подмножество', 'MSE', 'MAE', 'Max|err|', 'R²'],
    [
        ['Обучающее', '___', '___', '___', '___'],
        ['Контрольное', '___', '___', '___', '___'],
    ]
)
add_empty()

add_figure_placeholder('Рис. 6 — Аппроксимация и ошибка (Этап 2 — traingdx)')

doc.add_page_break()

# --- ЭТАП 3 ---
add_normal('Этап 3. Аппроксимация функции (метод второго порядка)', bold=True)
add_empty()

add_normal('Использована та же функция x = cos(t² − 2t + 3) и та же архитектура сети. '
           'Для обучения применён квазиньютоновский метод BFGS (trainbfg), '
           'реализованный через оптимизатор LBFGS.')
add_empty()

add_normal('Таблица 6 — Структура сети (Этап 3)')
add_table(
    ['Параметр', 'Значение'],
    [
        ['Архитектура', '[1] → [10, tansig] → [1, purelin]'],
        ['Алгоритм обучения', 'trainbfg (LBFGS — квазиньютоновский)'],
        ['Число эпох (epochs)', '600'],
        ['max_fail', '600'],
        ['goal', '1e-8'],
    ]
)
add_empty()

add_normal('Весовые коэффициенты и смещения (после обучения):', bold=True)
add_normal('Скрытый слой: W1 = [___], b1 = [___]')
add_normal('Выходной слой: W2 = [___], b2 = [___]')
add_empty()

add_figure_placeholder('Рис. 7 — Performance (Этап 3 — trainbfg)')
add_empty()

add_normal('Таблица 7 — Показатели качества (Этап 3)')
add_table(
    ['Подмножество', 'MSE', 'MAE', 'Max|err|', 'R²'],
    [
        ['Обучающее', '___', '___', '___', '___'],
        ['Контрольное', '___', '___', '___', '___'],
    ]
)
add_empty()

add_figure_placeholder('Рис. 8 — Аппроксимация и ошибка (Этап 3 — trainbfg)')
add_empty()

add_figure_placeholder('Рис. 9 — Сравнение методов обучения')

doc.add_page_break()

# ============================
# КОД ПРОГРАММЫ
# ============================
add_normal('Код программы', bold=True)
add_empty()

# Read code from notebook
import json
with open('/Users/arsknz/osnov/code/neural_networks/lab_2/main.ipynb') as f:
    nb = json.load(f)

for cell in nb['cells']:
    if cell['cell_type'] == 'code':
        src = ''.join(cell['source']) if isinstance(cell['source'], list) else cell['source']
        if src.strip():
            for line in src.split('\n'):
                add_code(line)
            add_empty()

doc.add_page_break()

# ============================
# ВЫВОД
# ============================
add_normal('Вывод', bold=True)
add_empty()

add_normal('В ходе лабораторной работы исследованы свойства многослойной нейронной сети '
           'прямого распространения и алгоритмы её обучения.')
add_empty()

add_normal('На первом этапе построена сеть для классификации трёх линейно неразделимых '
           'классов (две концентрические окружности и парабола). Сеть с архитектурой '
           '[2]→[20, tanh]→[3, sigmoid], обученная алгоритмом RProp, успешно '
           'классифицирует точки заданной области. Визуализация с помощью RGB-кодирования '
           'выходного вектора демонстрирует нелинейные границы разделения классов.')
add_empty()

add_normal('На втором этапе выполнена аппроксимация функции x = cos(t² − 2t + 3) '
           'с помощью метода первого порядка (traingdx — градиентный спуск с моментом '
           'и адаптивным шагом). Метод первого порядка сходится медленнее и требует '
           'большего числа эпох для достижения приемлемой точности.')
add_empty()

add_normal('На третьем этапе та же функция аппроксимирована с помощью метода '
           'второго порядка (trainbfg — квазиньютоновский метод BFGS/LBFGS). '
           'Метод второго порядка сходится значительно быстрее и достигает '
           'более низкой ошибки за меньшее число эпох, что подтверждает '
           'теоретическое преимущество методов второго порядка при обучении '
           'нейронных сетей.')
add_empty()

add_normal('Все этапы лабораторной работы выполнены. Графики и численные '
           'результаты соответствуют теоретическим ожиданиям.')

# Сохранение
out = '/Users/arsknz/osnov/code/neural_networks/lab_2/report_lab2.docx'
doc.save(out)
print(f'Отчёт сохранён: {out}')
